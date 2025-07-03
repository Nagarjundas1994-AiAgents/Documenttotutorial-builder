import os
import sys
import json
import re
from typing import TypedDict, List, Dict
from langgraph.graph import StateGraph, END
from openai import OpenAI
from openai.types.chat import ChatCompletionMessageParam
from dotenv import load_dotenv
import datetime

load_dotenv()

# --- Enhanced LLM Wrapper ---
class LLMWrapper:
    def __init__(self):
        # Try to use Google Gemini first, then XAI, then DeepSeek
        self.google_api_key = os.getenv('GOOGLE_API_KEY')
        self.xai_api_key = os.getenv('XAI_API_KEY')
        self.deepseek_api_key = os.getenv('DEEPSEEK_API_KEY')
        
        # Try Google first
        if self.google_api_key and self.google_api_key != "your-google-gemini-api-key":
            if self._setup_google():
                return
        
        # Try XAI second
        if self.xai_api_key and self.xai_api_key != "your-xai-api-key":
            if self._setup_xai():
                return
        
        # Try DeepSeek last
        if self.deepseek_api_key and self.deepseek_api_key != "your-deepseek-api-key":
            if self._setup_deepseek():
                return
        
        raise ValueError("No valid API key found for Google Gemini, XAI, or DeepSeek")
    
    def _setup_google(self):
        try:
            self.client = OpenAI(
                api_key=self.google_api_key, 
                base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
            )
            self.model_name = "gemini-2.0-flash"
            self.provider = "google"
            
            # Test the API with a simple call
            test_response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            print("✅ Using Google Gemini API")
            return True
        except Exception as e:
            print(f"❌ Google API failed: {str(e)[:100]}...")
            return False
    
    def _setup_xai(self):
        try:
            self.client = OpenAI(
                api_key=self.xai_api_key,
                base_url="https://api.x.ai/v1"
            )
            self.model_name = "grok-2-1212"
            self.provider = "xai"
            
            # Test the API with a simple call
            test_response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            print("✅ Using XAI (Grok) API")
            return True
        except Exception as e:
            print(f"❌ XAI API failed: {str(e)[:100]}...")
            return False
    
    def _setup_deepseek(self):
        try:
            self.client = OpenAI(
                api_key=self.deepseek_api_key,
                base_url="https://api.deepseek.com"
            )
            self.model_name = "deepseek-chat"
            self.provider = "deepseek"
            
            # Test the API with a simple call
            test_response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            print("✅ Using DeepSeek API")
            return True
        except Exception as e:
            print(f"❌ DeepSeek API failed: {str(e)[:100]}...")
            return False

    def invoke(self, prompt: str) -> str:
        # Customize system prompt based on provider
        if self.provider == "xai":
            system_content = "You are Grok, a highly intelligent, helpful AI assistant that creates clear, comprehensive tutorials with excellent explanations and practical examples."
        else:
            system_content = "You are a world-class technical writer and educator that creates clear, comprehensive tutorials with excellent explanations, practical examples, and engaging content."
        
        messages: List[ChatCompletionMessageParam] = [
            {"role": "system", "content": system_content},
            {"role": "user", "content": prompt}
        ]
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages,
                temperature=0.7,
                max_tokens=4000
            )
            content = response.choices[0].message.content
            return content if content is not None else ""
        except Exception as e:
            print(f"LLM API call failed: {e}")
            return f"Error generating content: {str(e)}"

# Initialize the LLM
llm = LLMWrapper()

# Define the enhanced state for our graph
class GraphState(TypedDict):
    original_query: str
    scraped_content: str
    tutorial_outline: Dict
    section_drafts: Dict[str, str]
    final_tutorial: str
    html_content: str
    error_message: str
    current_section_key: str
    enhanced_sections: Dict[str, str]
    code_examples: Dict[str, List[str]]
    concept_explanations: Dict[str, str]

# --- ENHANCED AGENT NODES ---

def extract_json_from_response(response: str) -> str:
    """Extract JSON from LLM response that might be wrapped in markdown code blocks."""
    if not response or response.strip() == "":
        raise ValueError("Empty response from LLM")
    
    # Try to find JSON in markdown code blocks first
    json_patterns = [
        r'```json\s*(\{.*?\})\s*```',  # ```json { ... } ```
        r'```\s*(\{.*?\})\s*```',      # ``` { ... } ```
        r'`(\{.*?\})`',                # `{ ... }`
    ]
    
    for pattern in json_patterns:
        json_match = re.search(pattern, response, re.DOTALL)
        if json_match:
            return json_match.group(1).strip()
    
    # Try to find JSON by looking for { and }
    start_idx = response.find('{')
    end_idx = response.rfind('}')
    
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        json_candidate = response[start_idx:end_idx+1]
        if json_candidate.strip().startswith('{') and json_candidate.strip().endswith('}'):
            return json_candidate.strip()
    
    return response.strip()

def generate_outline(state: GraphState) -> GraphState:
    """Generates a comprehensive, structured outline for the tutorial."""
    print("---AGENT: Generating Enhanced Outline---")
    prompt = f"""Based on the following documentation content, create a comprehensive, beginner-friendly tutorial outline that covers ALL important concepts without leaving anything behind.

Documentation Content:
---
{state['scraped_content'][:25000]}
---
User's original request: {state['original_query']}

Create a detailed JSON object with this exact structure:
{{
  "title": "Comprehensive Tutorial Title Here",
  "sections": [
    {{
      "title": "1. Introduction and Overview",
      "brief_description": "What this technology is, why it matters, and what learners will achieve"
    }},
    {{
      "title": "2. Prerequisites and Setup",
      "brief_description": "Required knowledge, tools, and environment setup"
    }},
    {{
      "title": "3. Core Concepts and Fundamentals",
      "brief_description": "Essential concepts that form the foundation"
    }},
    {{
      "title": "4. Getting Started - First Steps",
      "brief_description": "Hands-on introduction with simple examples"
    }},
    {{
      "title": "5. [Technology-Specific Section]",
      "brief_description": "Main features and capabilities"
    }},
    {{
      "title": "6. Advanced Features and Techniques",
      "brief_description": "More sophisticated usage patterns"
    }},
    {{
      "title": "7. Best Practices and Common Patterns",
      "brief_description": "Industry standards and recommended approaches"
    }},
    {{
      "title": "8. Troubleshooting and Common Issues",
      "brief_description": "Solutions to frequent problems"
    }},
    {{
      "title": "9. Real-World Examples and Use Cases",
      "brief_description": "Practical applications and complete examples"
    }},
    {{
      "title": "10. Next Steps and Advanced Resources",
      "brief_description": "Where to go from here and additional learning"
    }}
  ]
}}

IMPORTANT: 
- Ensure the outline covers EVERY important concept from the documentation
- Make sections comprehensive but digestible
- Include practical, hands-on elements
- Adapt section titles to match the specific technology/topic
- Aim for 8-12 sections total

Respond with ONLY the JSON object:"""

    try:
        outline_str = llm.invoke(prompt)
        print(f"Raw LLM response: {outline_str[:200]}...")
        
        json_str = extract_json_from_response(outline_str)
        print(f"Extracted JSON: {json_str[:200]}...")
        
        outline = json.loads(json_str)
        
        return {
            "original_query": state["original_query"],
            "scraped_content": state["scraped_content"],
            "tutorial_outline": outline,
            "section_drafts": {},
            "final_tutorial": "",
            "html_content": "",
            "error_message": "",
            "current_section_key": "0",
            "enhanced_sections": {},
            "code_examples": {},
            "concept_explanations": {}
        }
    except Exception as e:
        print(f"Outline generation error: {e}")
        return {
            "original_query": state["original_query"],
            "scraped_content": state["scraped_content"],
            "tutorial_outline": {},
            "section_drafts": {},
            "final_tutorial": "",
            "html_content": "",
            "error_message": f"Outline generation failed: {e}",
            "current_section_key": "0",
            "enhanced_sections": {},
            "code_examples": {},
            "concept_explanations": {}
        }

def write_enhanced_section(state: GraphState) -> GraphState:
    """Writes comprehensive, high-quality content for a single section."""
    print(f"---AGENT: Writing Enhanced Section: {state['current_section_key']}---")
    try:
        section_info = state['tutorial_outline']['sections'][int(state['current_section_key'])]
        
        # Generate comprehensive content
        main_prompt = f"""You are an expert technical writer creating a comprehensive tutorial section.
Your task is to write detailed, clear, and engaging content that leaves no concept unexplained.

CONTEXT:
Full Documentation: {state['scraped_content'][:20000]}
Tutorial Title: {state['tutorial_outline'].get('title', 'Tutorial')}
All Sections: {json.dumps([s['title'] for s in state['tutorial_outline'].get('sections', [])], indent=2)}

CURRENT SECTION TO WRITE:
- Section Title: {section_info['title']}
- Section Description: {section_info['brief_description']}

WRITING REQUIREMENTS:
1. **Comprehensive Coverage**: Cover ALL relevant concepts for this section
2. **Clear Explanations**: Use simple language, analogies, and examples
3. **Practical Examples**: Include working code examples where applicable
4. **Step-by-Step**: Break complex topics into digestible steps
5. **Visual Structure**: Use headers, lists, and formatting for readability
6. **Beginner-Friendly**: Assume no prior knowledge beyond prerequisites

CONTENT STRUCTURE:
- Start with a brief overview of what this section covers
- Explain concepts clearly with examples
- Include practical code examples (use ```language blocks)
- Add tips, warnings, or best practices where relevant
- End with a summary of key takeaways

FORMATTING GUIDELINES:
- Use ### for subsection headers
- Use **bold** for important terms
- Use `code` for inline code/commands
- Use ```language blocks for code examples
- Use > for important notes or tips
- Use numbered lists for procedures
- Use bullet points for features/concepts

Write comprehensive, detailed content for this section (aim for 800-1500 words):"""

        section_content = llm.invoke(main_prompt)
        
        # Enhance with key concepts
        concepts_prompt = f"""Extract the 3-5 most important concepts from this section and provide clear explanations:

Section: {section_info['title']}
Content: {section_content[:5000]}

For each concept, provide a clear, beginner-friendly explanation. Format as:
**Concept Name**: Clear explanation of what this is and why it matters.

Focus on the most essential concepts that beginners need to understand."""

        concepts_content = llm.invoke(concepts_prompt)
        
        # Generate practical examples
        examples_prompt = f"""Create 2-3 practical, working code examples for this section:

Section: {section_info['title']}
Content: {section_content[:5000]}

Each example should:
1. Be complete and runnable
2. Demonstrate key concepts
3. Include helpful comments
4. Show real-world usage

Format each as:
```language
// Example: Brief description
code here
```"""

        examples_content = llm.invoke(examples_prompt)
        
        # Combine all content
        enhanced_content = section_content
        
        if concepts_content and "**" in concepts_content:
            enhanced_content += f"\n\n### 🔑 Key Concepts\n\n{concepts_content}"
        
        if examples_content and "```" in examples_content:
            enhanced_content += f"\n\n### 💻 Practical Examples\n\n{examples_content}"
        
        # Add practice exercises
        exercise_prompt = f"""Create 1-2 simple practice exercises for this section:

Section: {section_info['title']}
Content: {section_content[:3000]}

Each exercise should:
1. Be achievable by beginners
2. Reinforce key concepts
3. Have clear instructions

Format as:
**Exercise 1**: Clear problem description
*Hint*: Helpful guidance
*Expected outcome*: What they should achieve"""

        exercise_content = llm.invoke(exercise_prompt)
        
        if exercise_content and "Exercise" in exercise_content:
            enhanced_content += f"\n\n### 🎯 Practice Exercises\n\n{exercise_content}"
        
        # Final enhancement pass
        enhanced_content = enhance_section_content(enhanced_content, section_info['title'])
        
        current_drafts = dict(state['section_drafts'])
        current_drafts[state['current_section_key']] = enhanced_content
        
        return {
            "original_query": state["original_query"],
            "scraped_content": state["scraped_content"],
            "tutorial_outline": state["tutorial_outline"],
            "section_drafts": current_drafts,
            "final_tutorial": state["final_tutorial"],
            "html_content": state.get("html_content", ""),
            "error_message": state["error_message"],
            "current_section_key": state["current_section_key"],
            "enhanced_sections": state.get("enhanced_sections", {}),
            "code_examples": state.get("code_examples", {}),
            "concept_explanations": state.get("concept_explanations", {})
        }
    except Exception as e:
        print(f"Section writing error: {e}")
        current_drafts = dict(state['section_drafts'])
        current_drafts[state['current_section_key']] = f"Error generating section content: {e}"
        
        return {
            "original_query": state["original_query"],
            "scraped_content": state["scraped_content"],
            "tutorial_outline": state["tutorial_outline"],
            "section_drafts": current_drafts,
            "final_tutorial": state["final_tutorial"],
            "html_content": state.get("html_content", ""),
            "error_message": state["error_message"],
            "current_section_key": state["current_section_key"],
            "enhanced_sections": state.get("enhanced_sections", {}),
            "code_examples": state.get("code_examples", {}),
            "concept_explanations": state.get("concept_explanations", {})
        }

def enhance_section_content(content: str, section_title: str) -> str:
    """Enhance section content with better formatting and structure."""
    # Add section introduction if missing
    if not content.strip().startswith(('##', '###', 'This section', 'In this section')):
        intro = f"## {section_title}\n\nThis section covers {section_title.lower()}. "
        content = intro + content
    
    # Ensure proper spacing around code blocks
    content = re.sub(r'(\n```)', r'\n\n```', content)
    content = re.sub(r'(```\n)', r'```\n\n', content)
    
    # Add proper spacing around headers
    content = re.sub(r'(\n)(#{1,6}\s)', r'\n\n\2', content)
    
    # Ensure proper list formatting
    content = re.sub(r'(\n)([0-9]+\.|\*|\-)\s', r'\n\n\2 ', content)
    
    return content.strip()

def create_beautiful_html(title: str, sections: list, metadata: dict) -> str:
    """Create a beautiful, elegant HTML document with modern styling."""
    
    # Generate table of contents
    toc_items = []
    for i, section in enumerate(sections):
        section_id = f"section-{i+1}"
        section_title = section.get("title", "Section").replace("## ", "").replace("# ", "")
        toc_items.append(f'<li><a href="#{section_id}" class="toc-link">{section_title}</a></li>')
    
    toc_html = f'<ul class="toc-list">{"".join(toc_items)}</ul>' if toc_items else ""
    
    # Generate sections HTML
    sections_html = []
    for i, section in enumerate(sections):
        section_id = f"section-{i+1}"
        section_title = section.get('title', 'Section').replace("## ", "").replace("# ", "")
        section_content = section.get('content', '')
        
        # Convert markdown-like content to HTML
        content_html = convert_markdown_to_html(section_content)
        
        sections_html.append(f'''
        <section id="{section_id}" class="tutorial-section">
            <h2 class="section-title">
                <span class="section-number">{i+1}</span>
                {section_title}
            </h2>
            <div class="section-content">
                {content_html}
            </div>
        </section>
        ''')
    
    # Create the complete HTML with enhanced styling
    html_template = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.8;
            color: #2d3748;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 25px 80px rgba(0,0,0,0.15);
            border-radius: 20px;
            overflow: hidden;
            margin-top: 2rem;
            margin-bottom: 2rem;
        }}
        
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 4rem 2rem;
            text-align: center;
            position: relative;
            overflow: hidden;
        }}
        
        .header::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: url('data:image/svg+xml,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100"><defs><pattern id="grain" width="100" height="100" patternUnits="userSpaceOnUse"><circle cx="25" cy="25" r="1" fill="white" opacity="0.1"/><circle cx="75" cy="75" r="1" fill="white" opacity="0.1"/><circle cx="50" cy="10" r="0.5" fill="white" opacity="0.1"/></pattern></defs><rect width="100" height="100" fill="url(%23grain)"/></svg>');
        }}
        
        .main-title {{
            font-size: 3rem;
            font-weight: 800;
            margin-bottom: 1rem;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
            position: relative;
            z-index: 1;
            letter-spacing: -0.02em;
        }}
        
        .subtitle {{
            font-size: 1.3rem;
            opacity: 0.9;
            font-weight: 300;
            position: relative;
            z-index: 1;
            max-width: 600px;
            margin: 0 auto;
        }}
        
        .metadata {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 2rem;
            border-left: 5px solid #667eea;
            margin: 0;
            font-size: 1rem;
            color: #495057;
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
        }}
        
        .metadata-item {{
            display: flex;
            align-items: center;
            margin: 0.5rem 0;
        }}
        
        .metadata-icon {{
            margin-right: 0.5rem;
            font-size: 1.2rem;
        }}
        
        .content {{
            padding: 3rem 2rem;
        }}
        
        .toc {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            border-radius: 20px;
            padding: 2.5rem;
            margin-bottom: 4rem;
            border: 1px solid #dee2e6;
            box-shadow: 0 10px 30px rgba(0,0,0,0.05);
        }}
        
        .toc h3 {{
            color: #495057;
            margin-bottom: 1.5rem;
            font-size: 1.5rem;
            display: flex;
            align-items: center;
            font-weight: 700;
        }}
        
        .toc h3::before {{
            content: "📚";
            margin-right: 0.75rem;
            font-size: 1.8rem;
        }}
        
        .toc-list {{
            list-style: none;
            padding: 0;
        }}
        
        .toc-list li {{
            margin: 0.75rem 0;
        }}
        
        .toc-link {{
            color: #495057;
            text-decoration: none;
            padding: 1rem 1.5rem;
            border-radius: 12px;
            display: block;
            transition: all 0.3s ease;
            border-left: 4px solid transparent;
            font-weight: 500;
            position: relative;
            overflow: hidden;
        }}
        
        .toc-link::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 0;
            height: 100%;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            transition: width 0.3s ease;
            z-index: -1;
        }}
        
        .toc-link:hover {{
            color: white;
            border-left-color: #667eea;
            transform: translateX(8px);
            box-shadow: 0 5px 20px rgba(102, 126, 234, 0.2);
        }}
        
        .toc-link:hover::before {{
            width: 100%;
        }}
        
        .tutorial-section {{
            margin-bottom: 5rem;
            scroll-margin-top: 2rem;
            position: relative;
        }}
        
        .section-title {{
            color: #2d3748;
            font-size: 2.2rem;
            margin-bottom: 2rem;
            padding-bottom: 1rem;
            border-bottom: 3px solid #667eea;
            display: flex;
            align-items: center;
            font-weight: 700;
            letter-spacing: -0.02em;
        }}
        
        .section-number {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            width: 3rem;
            height: 3rem;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 1.5rem;
            font-weight: bold;
            font-size: 1.2rem;
            box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3);
        }}
        
        .section-content {{
            font-size: 1.1rem;
            line-height: 1.8;
        }}
        
        .section-content p {{
            margin-bottom: 1.5rem;
            text-align: justify;
        }}
        
        .section-content h3 {{
            color: #4a5568;
            margin: 2.5rem 0 1.5rem 0;
            font-size: 1.5rem;
            font-weight: 600;
            border-left: 4px solid #667eea;
            padding-left: 1rem;
        }}
        
        .section-content h4 {{
            color: #718096;
            margin: 2rem 0 1rem 0;
            font-size: 1.2rem;
            font-weight: 600;
        }}
        
        .code-block {{
            background: linear-gradient(135deg, #2d3748 0%, #4a5568 100%);
            border: 1px solid #4a5568;
            border-radius: 15px;
            margin: 2rem 0;
            overflow: hidden;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
            position: relative;
        }}
        
        .code-block::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, #667eea, #764ba2, #667eea);
        }}
        
        .code-block pre {{
            padding: 2rem;
            margin: 0;
            overflow-x: auto;
            color: #e2e8f0;
            font-size: 0.95rem;
        }}
        
        .code-block code {{
            font-family: 'Fira Code', 'JetBrains Mono', 'Consolas', 'Monaco', monospace;
            line-height: 1.6;
        }}
        
        .inline-code {{
            background: linear-gradient(135deg, #f1f3f4 0%, #e8eaed 100%);
            color: #d63384;
            padding: 0.3rem 0.6rem;
            border-radius: 6px;
            font-family: 'Fira Code', 'Consolas', 'Monaco', monospace;
            font-size: 0.9em;
            font-weight: 500;
            border: 1px solid #dee2e6;
        }}
        
        .section-content ul {{
            margin: 1.5rem 0;
            padding-left: 0;
            list-style: none;
        }}
        
        .section-content li {{
            margin: 1rem 0;
            padding-left: 2rem;
            position: relative;
            line-height: 1.6;
        }}
        
        .section-content li::before {{
            content: "▶";
            color: #667eea;
            position: absolute;
            left: 0;
            font-weight: bold;
        }}
        
        .section-content ol {{
            margin: 1.5rem 0;
            padding-left: 2rem;
        }}
        
        .section-content ol li {{
            padding-left: 0.5rem;
        }}
        
        .section-content ol li::before {{
            display: none;
        }}
        
        .section-content strong {{
            color: #2d3748;
            font-weight: 700;
        }}
        
        .section-content em {{
            color: #4a5568;
            font-style: italic;
        }}
        
        .note-box {{
            background: linear-gradient(135deg, #e6fffa 0%, #b2f5ea 100%);
            border-left: 5px solid #38b2ac;
            padding: 1.5rem;
            margin: 2rem 0;
            border-radius: 0 10px 10px 0;
            box-shadow: 0 5px 15px rgba(56, 178, 172, 0.1);
        }}
        
        .note-box::before {{
            content: "💡 ";
            font-size: 1.2rem;
            margin-right: 0.5rem;
        }}
        
        .warning-box {{
            background: linear-gradient(135deg, #fef5e7 0%, #fbd38d 100%);
            border-left: 5px solid #ed8936;
            padding: 1.5rem;
            margin: 2rem 0;
            border-radius: 0 10px 10px 0;
            box-shadow: 0 5px 15px rgba(237, 137, 54, 0.1);
        }}
        
        .warning-box::before {{
            content: "⚠️ ";
            font-size: 1.2rem;
            margin-right: 0.5rem;
        }}
        
        .footer {{
            background: linear-gradient(135deg, #2d3748 0%, #4a5568 100%);
            color: white;
            text-align: center;
            padding: 3rem 2rem;
            margin-top: 4rem;
        }}
        
        .footer p {{
            margin: 0.5rem 0;
            opacity: 0.9;
        }}
        
        .footer-title {{
            font-size: 1.2rem;
            font-weight: 600;
            margin-bottom: 1rem;
        }}
        
        @media (max-width: 768px) {{
            .container {{
                margin: 1rem;
                border-radius: 15px;
            }}
            
            .header {{
                padding: 2rem 1rem;
            }}
            
            .main-title {{
                font-size: 2.2rem;
            }}
            
            .content {{
                padding: 2rem 1.5rem;
            }}
            
            .section-title {{
                font-size: 1.8rem;
                flex-direction: column;
                align-items: flex-start;
            }}
            
            .section-number {{
                margin-bottom: 1rem;
                margin-right: 0;
            }}
            
            .metadata {{
                flex-direction: column;
                align-items: flex-start;
            }}
        }}
        
        @media print {{
            body {{
                background: white;
            }}
            
            .container {{
                box-shadow: none;
                margin: 0;
            }}
            
            .header {{
                background: #667eea !important;
                -webkit-print-color-adjust: exact;
            }}
            
            .code-block {{
                background: #f8f9fa !important;
                color: #2d3748 !important;
                border: 1px solid #dee2e6 !important;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <header class="header">
            <h1 class="main-title">{title}</h1>
            <p class="subtitle">A comprehensive, step-by-step guide to mastering every concept</p>
        </header>
        
        <div class="metadata">
            <div class="metadata-item">
                <span class="metadata-icon">📅</span>
                <strong>Generated:</strong> {metadata.get('generated_at', 'Unknown')}
            </div>
            <div class="metadata-item">
                <span class="metadata-icon">🔗</span>
                <strong>Source:</strong> {metadata.get('source_query', 'Unknown')}
            </div>
            <div class="metadata-item">
                <span class="metadata-icon">📊</span>
                <strong>Sections:</strong> {len(sections)}
            </div>
        </div>
        
        <div class="content">
            <div class="toc">
                <h3>Table of Contents</h3>
                {toc_html}
            </div>
            
            {"".join(sections_html)}
        </div>
        
        <footer class="footer">
            <div class="footer-title">✨ Generated by Document to Tutorial Builder ✨</div>
            <p>Crafted with advanced AI for optimal learning experience</p>
            <p>Every concept explained • No detail left behind • Ready to learn</p>
        </footer>
    </div>
</body>
</html>'''
    
    return html_template

def convert_markdown_to_html(content: str) -> str:
    """Convert markdown-like content to HTML with enhanced formatting."""
    html_content = content
    
    # Convert code blocks with language detection
    html_content = re.sub(
        r'```(\w+)?\n(.*?)\n```',
        lambda m: f'<div class="code-block"><pre><code>{m.group(2).replace("<", "&lt;").replace(">", "&gt;")}</code></pre></div>',
        html_content,
        flags=re.DOTALL
    )
    
    # Convert inline code
    html_content = re.sub(r'`([^`]+)`', r'<code class="inline-code">\1</code>', html_content)
    
    # Convert bold text
    html_content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html_content)
    
    # Convert italic text
    html_content = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html_content)
    
    # Convert headers
    html_content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_content, flags=re.MULTILINE)
    html_content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_content, flags=re.MULTILINE)
    html_content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_content, flags=re.MULTILINE)
    
    # Convert notes and warnings
    html_content = re.sub(r'^> (.+)$', r'<div class="note-box">\1</div>', html_content, flags=re.MULTILINE)
    
    # Convert lists
    html_content = re.sub(r'^- (.+)$', r'<li>\1</li>', html_content, flags=re.MULTILINE)
    html_content = re.sub(r'^(\d+)\. (.+)$', r'<li>\2</li>', html_content, flags=re.MULTILINE)
    
    # Wrap consecutive list items in ul tags
    html_content = re.sub(r'(<li>.*?</li>)(?:\s*<li>.*?</li>)*', lambda m: f'<ul>{m.group(0)}</ul>', html_content, flags=re.DOTALL)
    
    # Convert line breaks to paragraphs
    paragraphs = html_content.split('\n\n')
    formatted_paragraphs = []
    
    for para in paragraphs:
        para = para.strip()
        if para and not para.startswith('<'):
            para = f'<p>{para.replace(chr(10), "<br>")}</p>'
        formatted_paragraphs.append(para)
    
    return '\n'.join(formatted_paragraphs)

def save_premium_formats(title: str, sections: list, metadata: dict, safe_title: str, timestamp: str, output_dir: str):
    """Save tutorial in 3 premium formats: HTML, PDF, DOCX."""
    saved_files = []
    
    # 1. Create Beautiful HTML
    try:
        html_content = create_beautiful_html(title, sections, metadata)
        
        html_filename = f"{safe_title}_{timestamp}.html"
        html_filepath = os.path.join(output_dir, html_filename)
        with open(html_filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        saved_files.append(("HTML", html_filepath, html_content))
        print(f"✅ Beautiful HTML saved: {html_filepath}")
    except Exception as e:
        print(f"❌ Error saving HTML: {e}")
        html_content = ""
    
    # 2. Create PDF with fallback
    try:
        pdf_filename = f"{safe_title}_{timestamp}.pdf"
        pdf_filepath = os.path.join(output_dir, pdf_filename)
        
        # Try WeasyPrint first
        try:
            import weasyprint
            
            # Create PDF with better styling for print
            pdf_html = html_content.replace(
                'background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);',
                'background: white;'
            )
            
            weasyprint.HTML(string=pdf_html).write_pdf(pdf_filepath)
            saved_files.append(("PDF", pdf_filepath, ""))
            print(f"✅ Professional PDF saved: {pdf_filepath}")
            
        except Exception as weasy_error:
            print(f"⚠️  WeasyPrint failed: {str(weasy_error)[:100]}...")
            print("🔄 Using alternative PDF generation...")
            
            # Use ReportLab as fallback
            from agents.pdf_fallback import create_pdf_with_reportlab
            
            if create_pdf_with_reportlab(title, sections, metadata, pdf_filepath):
                saved_files.append(("PDF", pdf_filepath, ""))
                print(f"✅ Alternative PDF saved: {pdf_filepath}")
            else:
                print(f"❌ PDF generation failed completely")
                
    except Exception as e:
        print(f"❌ Error in PDF generation: {e}")
    
    # 3. Create DOCX
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        
        # Set document margins
        sections_doc = doc.sections
        for section in sections_doc:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
        
        # Add title with styling
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(45, 55, 72)
        
        # Add subtitle
        subtitle_para = doc.add_paragraph("A comprehensive, step-by-step guide to mastering every concept")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(113, 128, 150)
        subtitle_run.italic = True
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Generated: {metadata.get('generated_at', 'Unknown')}").italic = True
        metadata_para.add_run("\n")
        metadata_para.add_run(f"Source: {metadata.get('source_query', 'Unknown')}").italic = True
        metadata_para.add_run("\n")
        metadata_para.add_run(f"Sections: {len(sections)}").italic = True
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table of contents
        toc_para = doc.add_heading("Table of Contents", level=1)
        toc_run = toc_para.runs[0]
        toc_run.font.color.rgb = RGBColor(102, 126, 234)
        
        for i, section_data in enumerate(sections):
            toc_item = doc.add_paragraph(f"{i+1}. {section_data.get('title', 'Section').replace('## ', '').replace('# ', '')}")
            toc_item.style = 'List Number'
        
        doc.add_page_break()
        
        # Add sections with enhanced formatting
        for i, section_data in enumerate(sections):
            # Section heading
            section_title = section_data.get('title', 'Section').replace('## ', '').replace('# ', '')
            section_heading = doc.add_heading(f"{i+1}. {section_title}", level=1)
            heading_run = section_heading.runs[0]
            heading_run.font.color.rgb = RGBColor(102, 126, 234)
            
            # Section content
            content = section_data.get('content', '')
            
            # Process content with better formatting
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Handle code blocks
                    if para.strip().startswith('```'):
                        code_content = re.sub(r'```.*?\n(.*?)\n```', r'\1', para, flags=re.DOTALL)
                        code_para = doc.add_paragraph(code_content)
                        code_para.style = 'Intense Quote'
                        # Style code paragraph
                        for run in code_para.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                    else:
                        # Clean up markdown formatting
                        para_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', para)  # Remove bold
                        para_clean = re.sub(r'\*([^*]+)\*', r'\1', para_clean)  # Remove italic
                        para_clean = re.sub(r'`([^`]+)`', r'\1', para_clean)  # Remove inline code
                        para_clean = re.sub(r'#{1,6}\s*', '', para_clean)  # Remove headers
                        para_clean = re.sub(r'^- ', '• ', para_clean, flags=re.MULTILINE)  # Convert bullets
                        
                        if para_clean.strip():
                            new_para = doc.add_paragraph(para_clean.strip())
                            # Add spacing
                            new_para.paragraph_format.space_after = Pt(6)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Save DOCX
        docx_filename = f"{safe_title}_{timestamp}.docx"
        docx_filepath = os.path.join(output_dir, docx_filename)
        doc.save(docx_filepath)
        saved_files.append(("DOCX", docx_filepath, ""))
        print(f"✅ Professional DOCX saved: {docx_filepath}")
        
    except Exception as e:
        print(f"❌ Error saving DOCX: {e}")
    
    return saved_files

def compile_tutorial(state: GraphState) -> GraphState:
    """Compiles all written sections into a final tutorial document and saves in premium formats only."""
    print("---AGENT: Compiling Final Tutorial---")
    outline = state['tutorial_outline']
    drafts = state['section_drafts']
    
    # Save to premium formats only (HTML, PDF, DOCX)
    try:
        # Create output directory if it doesn't exist
        output_dir = "generated_tutorials"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate safe filename
        title = outline.get('title', 'Comprehensive Tutorial')
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Prepare metadata
        metadata = {
            "title": title,
            "generated_at": datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "source_query": state.get('original_query', 'Unknown')
        }
        
        # Prepare sections data
        sections = [
            {
                "title": section.get('title', 'Section'),
                "brief_description": section.get('brief_description', ''),
                "content": drafts.get(str(i), "Error: Content for this section was not generated.")
            }
            for i, section in enumerate(outline.get('sections', []))
        ]
        
        # Save in premium formats only
        saved_files = save_premium_formats(title, sections, metadata, safe_title, timestamp, output_dir)
        
        print(f"\n🎉 Tutorial '{title}' saved in {len(saved_files)} premium formats:")
        for format_name, filepath, html_content in saved_files:
            print(f"   📄 {format_name}: {os.path.basename(filepath)}")
        
        # Store HTML content for frontend display
        html_content_for_frontend = ""
        for format_name, filepath, content in saved_files:
            if format_name == "HTML" and content:
                html_content_for_frontend = content
                break
        
        # Generate final markdown for compatibility
        final_md = f"# {title}\n\n"
        final_md += f"*Generated on: {metadata['generated_at']}*\n\n"
        final_md += f"*Source: {metadata['source_query']}*\n\n"
        final_md += "---\n\n"
        
        for i, section in enumerate(sections):
            final_md += f"## {i+1}. {section.get('title', 'Section')}\n\n"
            final_md += section.get('content', '') + "\n\n"
        
        return {
            "original_query": state.get("original_query", ""),
            "scraped_content": state.get("scraped_content", ""),
            "tutorial_outline": state.get("tutorial_outline", {}),
            "section_drafts": state.get("section_drafts", {}),
            "final_tutorial": final_md,
            "html_content": html_content_for_frontend,  # Add HTML for frontend
            "error_message": state.get("error_message", ""),
            "current_section_key": state.get("current_section_key", ""),
            "enhanced_sections": state.get("enhanced_sections", {}),
            "code_examples": state.get("code_examples", {}),
            "concept_explanations": state.get("concept_explanations", {})
        }
        
    except Exception as e:
        print(f"❌ Error saving tutorial: {e}")
        return {
            "original_query": state.get("original_query", ""),
            "scraped_content": state.get("scraped_content", ""),
            "tutorial_outline": state.get("tutorial_outline", {}),
            "section_drafts": state.get("section_drafts", {}),
            "final_tutorial": "",
            "html_content": "",
            "error_message": f"Tutorial compilation failed: {e}",
            "current_section_key": state.get("current_section_key", ""),
            "enhanced_sections": state.get("enhanced_sections", {}),
            "code_examples": state.get("code_examples", {}),
            "concept_explanations": state.get("concept_explanations", {})
        }

def should_continue(state: GraphState) -> str:
    """Determine whether to continue writing sections or compile the tutorial."""
    if state.get("error_message"):
        return "compile_tutorial"
    
    outline = state.get('tutorial_outline', {})
    sections = outline.get('sections', [])
    drafts = state.get('section_drafts', {})
    current_section_key = int(state.get('current_section_key', '0'))
    
    # Check if we have written all sections
    if len(drafts) >= len(sections):
        return "compile_tutorial"
    
    # Check if current section key is beyond available sections
    if current_section_key >= len(sections):
        return "compile_tutorial"
    
    return "write_section"

def get_next_section_key(state: GraphState) -> GraphState:
    """Updates the current section key for the next section to write."""
    outline = state.get('tutorial_outline', {})
    sections = outline.get('sections', [])
    drafts = state.get('section_drafts', {})
    
    # Find the next section that hasn't been written yet
    next_key = str(len(drafts))
    
    return {
        "original_query": state.get("original_query", ""),
        "scraped_content": state.get("scraped_content", ""),
        "tutorial_outline": state.get("tutorial_outline", {}),
        "section_drafts": state.get("section_drafts", {}),
        "final_tutorial": state.get("final_tutorial", ""),
        "html_content": state.get("html_content", ""),
        "error_message": state.get("error_message", ""),
        "current_section_key": next_key,
        "enhanced_sections": state.get("enhanced_sections", {}),
        "code_examples": state.get("code_examples", {}),
        "concept_explanations": state.get("concept_explanations", {})
    }

def create_tutorial_graph():
    """Create the enhanced tutorial generation graph."""
    workflow = StateGraph(GraphState)
    workflow.add_node("generate_outline", generate_outline)
    workflow.add_node("write_section", write_enhanced_section)
    workflow.add_node("compile_tutorial", compile_tutorial)
    workflow.add_node("get_next_section_key", get_next_section_key)
    
    workflow.set_entry_point("generate_outline")
    
    workflow.add_conditional_edges(
        "generate_outline",
        lambda state: "compile_tutorial" if state.get("error_message") else "get_next_section_key",
        {
            "get_next_section_key": "get_next_section_key",
            "compile_tutorial": "compile_tutorial"
        }
    )
    
    workflow.add_edge("get_next_section_key", "write_section")
    
    workflow.add_conditional_edges(
        "write_section",
        should_continue,
        {
            "write_section": "get_next_section_key",
            "compile_tutorial": "compile_tutorial"
        }
    )
    
    workflow.add_edge("compile_tutorial", END)
    
    app = workflow.compile()
    return app